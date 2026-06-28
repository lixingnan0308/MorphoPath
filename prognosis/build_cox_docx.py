#!/usr/bin/env python3
"""Word (.docx) version of the multivariate Cox table, mirroring the booktabs
PDF: bold variable sub-headers (Model probability / Age / grade / Sex), cohort
rows (BTH/FMUUH/TCGA) beneath, columns Cohort | n | events | HR (95% CI) | P,
booktabs rules (header top sz8, band tops sz4, last bottom sz8), title and a
small italic model note.
"""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
CSV = ROOT / "results" / "multivariate_cox_sex.csv"
OUT = ROOT / "results" / "multivariate_cox_table.docx"

LABEL = {"prob": "Model probability", "Age": "Age", "grade_high": "grade", "sex_male": "Sex"}
VAR_ORDER = ["prob", "Age", "grade_high", "sex_male"]
COHORTS = ["BTH", "FMUUH", "TCGA"]
HEADERS = ["Cohort", "n", "events", "HR (95% CI)", "P"]
BODY_PT = 11
NOTE = ("Model: Hazard ~ Model probability + Age + grade + Sex (slide level; "
        "cohort-specific fits). HR for Sex is Male vs. Female; grade is WHO grade ≥3 vs. <3.")


def fmt_hr(r):
    return f"{r.HR:.2f} ({r.lo95:.2f}–{r.hi95:.2f})"


def fmt_p(p):
    return "<0.001" if p < 0.001 else f"{p:.3f}"


def edge(name, val, sz=None):
    e = OxmlElement(f"w:{name}")
    e.set(qn("w:val"), val)
    if sz is not None:
        e.set(qn("w:sz"), str(sz)); e.set(qn("w:space"), "0"); e.set(qn("w:color"), "auto")
    return e


def set_borders(cell, top=("nil", None), bottom=("nil", None)):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    b = OxmlElement("w:tcBorders")
    b.append(edge("top", *top)); b.append(edge("left", "nil"))
    b.append(edge("bottom", *bottom)); b.append(edge("right", "nil"))
    tcPr.append(b)


def cell_text(cell, text, bold=False, size=BODY_PT, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold


res = pd.read_csv(CSV)
doc = Document()

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = t.add_run("Multivariate Cox analysis")
tr.font.bold = True
tr.font.size = Pt(14)

# Build row plan: header + (band + 3 cohorts) x 4
plan = [("header", None)]
for v in VAR_ORDER:
    plan.append(("band", v))
    for c in COHORTS:
        plan.append(("row", (v, c)))
n_rows = len(plan)

tbl = doc.add_table(rows=n_rows, cols=len(HEADERS))
tbl.style = "Normal Table"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

last = n_rows - 1
for ri, (kind, payload) in enumerate(plan):
    row = tbl.rows[ri]
    top = ("nil", None)
    bottom = ("nil", None)
    if ri == 0:
        top = ("single", 8)            # toprule
    elif kind == "band":
        top = ("single", 4)            # midrule before each variable block
    if ri == last:
        bottom = ("single", 8)         # bottomrule

    if kind == "header":
        for j, h in enumerate(HEADERS):
            cell_text(row.cells[j], h, bold=False, align="left" if j == 0 else "center")
    elif kind == "band":
        merged = row.cells[0]
        for j in range(1, len(HEADERS)):
            merged = merged.merge(row.cells[j])
        cell_text(merged, LABEL[payload], bold=True, align="left")
    else:
        v, c = payload
        r = res[(res.variable == v) & (res.cohort == c)].iloc[0]
        vals = [f" {c}", str(int(r.n)), str(int(r.events)), fmt_hr(r), fmt_p(r.P)]
        for j, val in enumerate(vals):
            cell_text(row.cells[j], val, bold=False, align="left" if j == 0 else "center")

    for ccell in row.cells:
        set_borders(ccell, top, bottom)

# Note
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(8)
nr = note.add_run(NOTE)
nr.font.italic = True
nr.font.size = Pt(9)

doc.save(OUT)
print(f"wrote {OUT}")
print(f"rows={n_rows} (header + 4 vars x (band+3 cohorts))")
